"""
Automated test runner for the fairness-aware movie recommender.
Runs predefined multi-turn conversations and writes results to a Word document.

Usage (from inside the agent/ folder):
    python ../run_tests.py

Requirements:
    pip install python-docx python-dotenv
    (all agent deps must also be installed)
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "agent"))

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

from graph import build_graph, DEFAULT_DB_SCHEMA

# ── Test scenarios ─────────────────────────────────────────────────────────────
# Each scenario is a dict with:
#   name  : section heading
#   turns : list of user messages in order
SCENARIOS = [
    {
        "name": "Director gender — action movies",
        "turns": [
            "Give me some popular action movies",
            "Why are all of these directed by men?",
        ],
    },
    {
        "name": "Director gender — drama movies",
        "turns": [
            "Recommend me some acclaimed dramas",
            "Are there any female directors in these results?",
        ],
    },
    {
        "name": "Origin country — thriller movies",
        "turns": [
            "Give me 4 good thriller movies",
            "Why are all of these American films?",
        ],
    },
    {
        "name": "Origin country — Asian cinema",
        "turns": [
            "Show me 7 highly rated dramas",
            "Why is there no representation from Asian cinema here?",
        ],
    },
    {
        "name": "Genre diversity",
        "turns": [
            "Give me 9 good movies to watch",
            "These are all action movies, where's the variety?",
        ],
    },
    {
        "name": "Repeated fairness challenge — non-English",
        "turns": [
            "Give me 3 great movies",
            "Why are all these English movies?",
            "There is still 1 English movie, add more non-English",
        ],
    },
    {
        "name": "Director gender — small result set (k=3)",
        "turns": [
            "Give me the 6 best action movies",
            "Why are they all by male directors?",
        ],
    },
]


# ── Graph runner ───────────────────────────────────────────────────────────────

def run_scenario(graph, scenario: dict) -> list[dict]:
    """Run a multi-turn scenario and return a list of turn result dicts."""
    state = {
        "messages": [],
        "agent_output": None,
        "route": None,
        "db_schema": DEFAULT_DB_SCHEMA,
        "error": None,
        "query_result": None,
        "last_results": None,
        "last_query": None,
        "fairness_report": None,
        "vector_search_query": None,
        "vector_embeddings": None,
        "use_vector_search": False,
        "latest_user_question": None,
        "sql_error": None,
        "sql_correction_attempted": False,
    }

    turns = []
    for user_text in scenario["turns"]:
        state["messages"].append({"role": "user", "content": user_text})
        print(f"    > {user_text[:80]}")

        result = graph.invoke(state)
        state = result

        output = result.get("agent_output")
        fairness_report = result.get("fairness_report")

        # Extract display text
        if isinstance(output, dict):
            display_text = output.get("explanation", "")
            reranked = output.get("results", [])
        elif output:
            display_text = output
            reranked = result.get("query_result") or []
        else:
            display_text = result.get("error") or "No output."
            reranked = []

        # Append to message history
        if display_text:
            state["messages"].append({"role": "assistant", "content": display_text})

        turns.append({
            "user": user_text,
            "assistant": display_text,
            "results": reranked,
            "fairness_report": fairness_report,
            "error": result.get("error"),
        })

    return turns


# ── Word document builder ──────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_metric_table(doc, before: dict, after: dict):
    """Add a before/after fairness metrics table."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Set column widths (total ~4.5 inches)
    widths = [Inches(1.5), Inches(1.5), Inches(1.5)]
    for i, cell in enumerate(table.columns[0].cells):
        cell.width = widths[0]

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Metric", "Before", "After"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def fmt(v):
        if v is None:
            return "N/A"
        try:
            return f"{float(v):+.4f}"
        except Exception:
            return str(v)

    exp_b = before.get("exposure_k") or {}
    exp_a = after.get("exposure_k") or {}

    rows = [
        ("SPD",       before.get("spd"),                   after.get("spd")),
        ("EOD",       before.get("eod"),                   after.get("eod")),
        ("OAED",      before.get("oaed"),                  after.get("oaed")),
        ("Exp@K (G+)", exp_b.get("protected"),             exp_a.get("protected")),
        ("Exp@K (G-)", exp_b.get("unprotected"),           exp_a.get("unprotected")),
        ("G+ count",  before.get("protected_count"),       after.get("protected_count")),
        ("G- count",  before.get("unprotected_count"),     after.get("unprotected_count")),
    ]

    for name, b, a in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = fmt(b)
        row[2].text = fmt(a)
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_results_table(doc, results: list[dict], label: str):
    """Add a movie results table."""
    if not results:
        doc.add_paragraph("No results returned.", style="Normal")
        return

    doc.add_paragraph(label, style="Normal").runs[0].bold = True
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["#", "Title", "Year", "★"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True

    for i, r in enumerate(results, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = r.get("title") or "Unknown"
        row[2].text = str(r.get("release_year") or "")
        rating = r.get("vote_average") or r.get("ratings") or r.get("rating")
        row[3].text = f"{float(rating):.2f}" if rating else "N/A"


def build_document(all_results: list[dict]) -> Document:
    doc = Document()

    # Title
    title = doc.add_heading("Fairness-Aware Movie Recommender — Test Results", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    ts = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for scenario_result in all_results:
        name = scenario_result["name"]
        turns = scenario_result["turns"]

        add_heading(doc, name, level=1)

        for t, turn in enumerate(turns, 1):
            add_heading(doc, f"Turn {t}", level=2)

            # User message
            p = doc.add_paragraph()
            p.add_run("You: ").bold = True
            p.add_run(turn["user"])

            # Assistant response
            p = doc.add_paragraph()
            p.add_run("Assistant: ").bold = True
            p.add_run(turn["assistant"] or "(no output)")

            # Initial results (turn 1 only)
            if t == 1 and turn["results"]:
                add_results_table(doc, turn["results"], "Recommendations:")

            # Fairness report
            report = turn.get("fairness_report")
            if report:
                attr       = report.get("attribute", "")
                p_lbl      = report.get("protected_label", "")
                u_lbl      = report.get("unprotected_label", "")
                before     = report.get("metrics_before") or {}
                after      = report.get("metrics_after") or {}
                reranked   = report.get("reranked_results") or []
                supp_count = report.get("supplementary_count", 0)

                doc.add_paragraph("")
                p = doc.add_paragraph()
                p.add_run("Fairness assessment").bold = True

                meta = doc.add_paragraph()
                meta.add_run(f"Attribute: ").bold = True
                meta.add_run(f"{attr}   ")
                meta.add_run(f"G+: ").bold = True
                meta.add_run(f"{p_lbl}   ")
                meta.add_run(f"G−: ").bold = True
                meta.add_run(u_lbl)

                p = doc.add_paragraph()
                p.add_run(f"Films supplemented: {supp_count}")

                add_metric_table(doc, before, after)
                doc.add_paragraph("")

                if reranked:
                    add_results_table(doc, reranked, "Re-ranked results:")

            if turn.get("error"):
                p = doc.add_paragraph()
                p.add_run("Error: ").bold = True
                run = p.add_run(turn["error"])
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            doc.add_paragraph("")

        doc.add_page_break()

    return doc


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    print("Building graph...")
    graph = build_graph()

    all_results = []
    for i, scenario in enumerate(SCENARIOS, 1):
        print(f"\n[{i}/{len(SCENARIOS)}] {scenario['name']}")
        turns = run_scenario(graph, scenario)
        all_results.append({"name": scenario["name"], "turns": turns})
        print(f"    done — {len(turns)} turns")

    print("\nBuilding Word document...")
    doc = build_document(all_results)

    out_path = os.path.join(os.path.dirname(__file__), "..", "test_results.docx")
    doc.save(out_path)
    print(f"Saved → {out_path}")


if __name__ == "__main__":
    main()